"""
DOCX text extraction utilities using python-docx.
Extracts text from paragraphs and tables with placeholder handling.
"""

import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Generator, Union

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from ..logging_setup import get_logger

logger = get_logger("extractors.docx")


def iter_block_items(parent: Union[DocumentType, _Cell]) -> Generator[Union[Paragraph, Table], None, None]:
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    
    Args:
        parent: Document or Cell object to iterate through
        
    Yields:
        Paragraph or Table objects in document order
    """
    if isinstance(parent, DocumentType):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent must be Document or _Cell instance")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_paragraph_text_with_fields(paragraph: Paragraph) -> str:
    """
    Extract text from a paragraph, including MERGEFIELD placeholders.
    
    Args:
        paragraph: Paragraph object from python-docx
        
    Returns:
        Complete text content of the paragraph
    """
    text = ""
    for run in paragraph.runs:
        text += run.text
    return text


def extract_content(docx_path: str) -> str:
    """
    Extract document content using python-docx library with proper table handling.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content with placeholders normalized
        
    Raises:
        Exception: If document cannot be opened or processed
    """
    logger.debug(f"Starting text extraction from: {docx_path}")
    
    try:
        # Try to open the document with error handling for custom XML issues
        try:
            document = Document(docx_path)
        except Exception as e:
            error_msg = str(e)
            if "customXML" in error_msg and "archive" in error_msg:
                logger.warning(f"Document has corrupted custom XML parts, attempting alternative extraction: {docx_path}")
                # Try to extract using zipfile directly as a fallback
                return _extract_with_zipfile_fallback(docx_path)
            else:
                raise e
        
        all_text = []
        
        # Extract headers first
        for section in document.sections:
            if section.header.is_linked_to_previous:
                continue  # Skip if linked to previous section header
            
            for paragraph in section.header.paragraphs:
                header_text = get_paragraph_text_with_fields(paragraph).strip()
                if header_text:
                    all_text.append(header_text)
            
            # Check for header tables
            for table in section.header.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            cell_text += get_paragraph_text_with_fields(paragraph)
                        row_text.append(cell_text.strip())
                    
                    full_row_text = "\t".join(row_text).strip()
                    if full_row_text:
                        all_text.append(full_row_text)
        
        # Add separator after headers if any were found
        if any(text for text in all_text):
            all_text.append("")
        
        # Regular expression to match placeholder braces and normalize them
        # placeholder_re = re.compile(r"\{\s*([^{}\s].*?)\s*\}")
        
        for block in iter_block_items(document):
            if isinstance(block, Paragraph):
                paragraph_text = get_paragraph_text_with_fields(block).strip()
                if paragraph_text:
                    # Replace placeholder braces with just the placeholder name
                    # paragraph_text = placeholder_re.sub(r'\1', paragraph_text)
                    all_text.append(paragraph_text)
                    
            elif isinstance(block, Table):
                for row in block.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            cell_text += get_paragraph_text_with_fields(paragraph)
                        # Replace placeholder braces with just the placeholder name
                        # cell_text = placeholder_re.sub(r'\1', cell_text).strip()
                        row_text.append(cell_text)
                    
                    # Join cell text with a tab to represent table columns
                    full_row_text = "\t".join(row_text).strip()
                    if full_row_text:
                        all_text.append(full_row_text)
        
        # Extract footers last
        footer_added = False
        for section in document.sections:
            if section.footer.is_linked_to_previous:
                continue  # Skip if linked to previous section footer
            
            # Add separator before first footer
            if not footer_added:
                all_text.append("")
                footer_added = True
            
            for paragraph in section.footer.paragraphs:
                footer_text = get_paragraph_text_with_fields(paragraph).strip()
                if footer_text:
                    all_text.append(footer_text)
            
            # Check for footer tables
            for table in section.footer.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            cell_text += get_paragraph_text_with_fields(paragraph)
                        row_text.append(cell_text.strip())
                    
                    full_row_text = "\t".join(row_text).strip()
                    if full_row_text:
                        all_text.append(full_row_text)
        
        extracted_text = '\n'.join(all_text)
        
        # Count headers and footers for logging
        header_count = sum(1 for section in document.sections if not section.header.is_linked_to_previous and 
                          (any(p.text.strip() for p in section.header.paragraphs) or section.header.tables))
        footer_count = sum(1 for section in document.sections if not section.footer.is_linked_to_previous and 
                          (any(p.text.strip() for p in section.footer.paragraphs) or section.footer.tables))
        
        parts_info = ["main document"]
        if header_count > 0:
            parts_info.insert(0, f"{header_count} header(s)")
        if footer_count > 0:
            parts_info.append(f"{footer_count} footer(s)")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from {', '.join(parts_info)} from {docx_path}")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Failed to extract content from {docx_path}: {e}")
        raise


def _extract_with_zipfile_fallback(docx_path: str) -> str:
    """
    Fallback extraction method for DOCX files with corrupted custom XML.
    Uses direct XML parsing of document.xml, headers, and footers.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content including headers and footers
    """
    logger = get_logger(__name__)
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            all_content = []
            
            # DOCX uses the WordProcessingML namespace
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            def extract_text_from_xml(xml_content, part_name):
                """Extract text from XML content"""
                try:
                    tree = ET.parse(xml_content)
                    root = tree.getroot()
                    paragraphs = []
                    
                    # Process each paragraph separately to preserve structure
                    for para_elem in root.findall('.//w:p', namespaces):
                        para_text_parts = []
                        
                        # Get all text runs within this paragraph
                        for text_elem in para_elem.findall('.//w:t', namespaces):
                            if text_elem.text:
                                para_text_parts.append(text_elem.text)
                        
                        # Join text runs within paragraph and add to paragraphs
                        if para_text_parts:
                            paragraph_text = ''.join(para_text_parts).strip()
                            if paragraph_text:  # Only add non-empty paragraphs
                                paragraphs.append(paragraph_text)
                    
                    # Also extract table content if present
                    for table_elem in root.findall('.//w:tbl', namespaces):
                        for row_elem in table_elem.findall('.//w:tr', namespaces):
                            cell_texts = []
                            for cell_elem in row_elem.findall('.//w:tc', namespaces):
                                cell_text_parts = []
                                for text_elem in cell_elem.findall('.//w:t', namespaces):
                                    if text_elem.text:
                                        cell_text_parts.append(text_elem.text)
                                if cell_text_parts:
                                    cell_texts.append(''.join(cell_text_parts).strip())
                            
                            if cell_texts:
                                # Join cells with tabs, add row to paragraphs
                                row_text = '\t'.join(cell_texts)
                                if row_text.strip():
                                    paragraphs.append(row_text)
                    
                    if paragraphs:
                        logger.debug(f"Extracted {len(paragraphs)} paragraphs from {part_name}")
                        return paragraphs
                    else:
                        return []
                        
                except Exception as e:
                    logger.warning(f"Failed to extract from {part_name}: {e}")
                    return []
            
            # Extract headers first (usually appear at top of document)
            file_list = zip_file.namelist()
            header_files = [f for f in file_list if f.startswith('word/header') and f.endswith('.xml')]
            header_files.sort()  # Process in order
            
            for header_file in header_files:
                try:
                    with zip_file.open(header_file) as header_xml:
                        header_content = extract_text_from_xml(header_xml, header_file)
                        if header_content:
                            all_content.extend(header_content)
                            all_content.append("")  # Add separator line after header
                except Exception as e:
                    logger.warning(f"Failed to process {header_file}: {e}")
            
            # Extract main document content
            try:
                with zip_file.open('word/document.xml') as doc_xml:
                    doc_content = extract_text_from_xml(doc_xml, 'word/document.xml')
                    if doc_content:
                        all_content.extend(doc_content)
                        
            except KeyError:
                logger.error(f"Could not find word/document.xml in {docx_path}")
                return ""
            
            # Extract footers last (usually appear at bottom of document)
            footer_files = [f for f in file_list if f.startswith('word/footer') and f.endswith('.xml')]
            footer_files.sort()  # Process in order
            
            if footer_files:
                all_content.append("")  # Add separator line before footer
                
            for footer_file in footer_files:
                try:
                    with zip_file.open(footer_file) as footer_xml:
                        footer_content = extract_text_from_xml(footer_xml, footer_file)
                        if footer_content:
                            all_content.extend(footer_content)
                except Exception as e:
                    logger.warning(f"Failed to process {footer_file}: {e}")
            
            # Join all content
            extracted_text = '\n'.join(all_content).strip()
            
            if extracted_text:
                parts_found = []
                if header_files:
                    parts_found.append(f"{len(header_files)} header(s)")
                parts_found.append("main document")
                if footer_files:
                    parts_found.append(f"{len(footer_files)} footer(s)")
                
                logger.info(f"Successfully extracted {len(extracted_text)} characters from {', '.join(parts_found)} using fallback method from {docx_path}")
                return extracted_text
            else:
                logger.warning(f"No text content found in any document parts for {docx_path}")
                return ""
                
    except Exception as e:
        logger.error(f"Fallback extraction failed for {docx_path}: {e}")
        return ""
